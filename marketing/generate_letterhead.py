import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch, mm
from reportlab.lib.colors import HexColor
from reportlab.graphics import renderPDF
from svglib.svglib import svg2rlg
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics
from reportlab.graphics import renderPM

def create_pdf(logo_path, output_path):
    c = canvas.Canvas(output_path, pagesize=A4)
    width, height = A4

    # Setup Colors
    GOLD = HexColor("#D4AF6A")
    DARK = HexColor("#080706")
    GREY = HexColor("#555555")

    # Add Logo
    drawing = svg2rlg(logo_path)
    
    # Scale logo appropriately
    # Original might be large, scale to fit 2 inches wide
    scale_factor = (2 * inch) / drawing.width
    drawing.width = drawing.width * scale_factor
    drawing.height = drawing.height * scale_factor
    drawing.scale(scale_factor, scale_factor)

    # Position logo at top left (x=20mm, y=A4_height - 30mm)
    renderPDF.draw(drawing, c, 20*mm, height - 30*mm)

    # Add Company Details (Top Right)
    c.setFont("Helvetica-Bold", 10)
    c.setFillColor(GOLD)
    c.drawRightString(width - 20*mm, height - 20*mm, "BIGWEB DIGITAL LIMITED")
    
    c.setFont("Helvetica", 9)
    c.setFillColor(DARK)
    c.drawRightString(width - 20*mm, height - 25*mm, "124 Elite Avenue, Tech District")
    c.drawRightString(width - 20*mm, height - 30*mm, "hello@bigweb.digital | www.bigweb.digital")
    c.drawRightString(width - 20*mm, height - 35*mm, "+234 (0) 800 000 0000")

    # Draw gold accent line
    c.setStrokeColor(GOLD)
    c.setLineWidth(1.5)
    c.line(20*mm, height - 40*mm, width - 20*mm, height - 40*mm)

    # Add placeholder content
    c.setFont("Helvetica", 11)
    c.drawString(20*mm, height - 60*mm, "Date: [Date]")
    c.drawString(20*mm, height - 66*mm, "Ref: [Reference]")

    c.setFont("Helvetica-Bold", 11)
    c.drawString(20*mm, height - 80*mm, "Tobi Adeyemi")
    c.setFont("Helvetica", 11)
    c.drawString(20*mm, height - 86*mm, "Founder & Lead Strategist")
    c.drawString(20*mm, height - 92*mm, "[Company Name]")
    c.drawString(20*mm, height - 98*mm, "[Address Line 1]")
    c.drawString(20*mm, height - 104*mm, "[Address Line 2]")

    c.drawString(20*mm, height - 120*mm, "Dear [Name],")
    
    text = c.beginText(20*mm, height - 135*mm)
    text.setFont("Helvetica", 11)
    text.setLeading(16)
    text.textLines("""[Begin your highly professional and impactful letter here. The typography is set with 
a comfortable line height for optimal readability and a premium feel. The margins are 
generously proportioned to frame the text elegantly.]

[This is a secondary paragraph. Continue your text. The gold accent line at the top 
separates the branding from the content cleanly.]""")
    c.drawText(text)

    c.drawString(20*mm, height - 180*mm, "Sincerely,")
    
    c.setFont("Helvetica-Bold", 11)
    c.drawString(20*mm, height - 210*mm, "DANIEL ORIAZOWAN")
    c.setFont("Helvetica", 11)
    c.drawString(20*mm, height - 216*mm, "FOUNDER & LEAD FRONTEND ENGINEER")
    c.drawString(20*mm, height - 222*mm, "BIGWEB Digital Limited")

    # Add Footer
    c.setStrokeColor(HexColor("#E5E5E5"))
    c.setLineWidth(1)
    c.line(20*mm, 25*mm, width - 20*mm, 25*mm)
    
    c.setFont("Helvetica", 8)
    c.setFillColor(GREY)
    
    footer_y = 18*mm
    # Center text manually using stringWidth
    part1 = "BIGWEB Digital Limited  |  "
    part2 = "CAC RC No: 8730433"
    part3 = "  |  Transforming digital potential into measurable revenue."
    
    w1 = c.stringWidth(part1, "Helvetica", 8)
    w2 = c.stringWidth(part2, "Helvetica-Bold", 8)
    w3 = c.stringWidth(part3, "Helvetica", 8)
    
    total_w = w1 + w2 + w3
    start_x = (width - total_w) / 2
    
    c.drawString(start_x, footer_y, part1)
    c.setFont("Helvetica-Bold", 8)
    c.setFillColor(GOLD)
    c.drawString(start_x + w1, footer_y, part2)
    c.setFont("Helvetica", 8)
    c.setFillColor(GREY)
    c.drawString(start_x + w1 + w2, footer_y, part3)

    c.save()
    print(f"Generated PDF: {output_path}")

def create_docx(logo_path, output_path):
    doc = Document()

    # Convert SVG to PNG for DOCX
    png_path = "temp_logo.png"
    drawing = svg2rlg(logo_path)
    renderPM.drawToFile(drawing, png_path, fmt="PNG", dpi=300)

    # Setup styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.color.rgb = RGBColor(8, 7, 6) # Dark Charcoal

    # Header section (Using a 2-column table for layout)
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(3)
    table.columns[1].width = Inches(3)

    cell1 = table.cell(0, 0)
    p1 = cell1.paragraphs[0]
    run = p1.add_run()
    run.add_picture(png_path, width=Inches(2.0))

    cell2 = table.cell(0, 1)
    p2 = cell2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    r1 = p2.add_run("BIGWEB DIGITAL LIMITED\n")
    r1.bold = True
    r1.font.color.rgb = RGBColor(212, 175, 106) # Gold
    
    p2.add_run("124 Elite Avenue, Tech District\n")
    p2.add_run("hello@bigweb.digital | www.bigweb.digital\n")
    p2.add_run("+234 (0) 800 000 0000")

    doc.add_paragraph() # spacing

    # Content
    doc.add_paragraph("Date: [Date]")
    doc.add_paragraph("Ref: [Reference]\n")

    p = doc.add_paragraph()
    p.add_run("Tobi Adeyemi").bold = True
    doc.add_paragraph("Founder & Lead Strategist\n[Company Name]\n[Address Line 1]\n[Address Line 2]\n")

    doc.add_paragraph("Dear [Name],")
    
    p = doc.add_paragraph("[Begin your highly professional and impactful letter here. The typography is set with a comfortable line height for optimal readability and a premium feel.]")
    p.paragraph_format.space_after = Pt(12)
    
    p = doc.add_paragraph("[This is a secondary paragraph. Continue your text. The letterhead balances brand presence with minimal distractions.]")
    p.paragraph_format.space_after = Pt(24)

    doc.add_paragraph("Sincerely,\n\n")

    p = doc.add_paragraph()
    p.add_run("DANIEL ORIAZOWAN").bold = True
    doc.add_paragraph("FOUNDER & LEAD FRONTEND ENGINEER\nBIGWEB Digital Limited")

    # Add Footer
    section = doc.sections[0]
    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    footer_p.add_run("BIGWEB Digital Limited  |  ").font.color.rgb = RGBColor(85, 85, 85)
    r2 = footer_p.add_run("CAC RC No: 8730433")
    r2.bold = True
    r2.font.color.rgb = RGBColor(212, 175, 106)
    footer_p.add_run("  |  Transforming digital potential into measurable revenue.").font.color.rgb = RGBColor(85, 85, 85)

    doc.save(output_path)
    
    # Cleanup temp PNG
    if os.path.exists(png_path):
        os.remove(png_path)
    
    print(f"Generated DOCX: {output_path}")

if __name__ == "__main__":
    logo = os.path.join(os.path.dirname(__file__), "..", "logo", "05_logo_horizontal_light_bg.svg")
    pdf_out = os.path.join(os.path.dirname(__file__), "BIGWEB_Letterhead.pdf")
    docx_out = os.path.join(os.path.dirname(__file__), "BIGWEB_Letterhead.docx")
    
    create_pdf(logo, pdf_out)
    create_docx(logo, docx_out)
    print("Done!")
